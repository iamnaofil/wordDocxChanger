from docx import Document

def replace_placeholder(doc, placeholder, value):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(placeholder, value)

def main():
    template_path = "/content/Cover Letter.docx"  # Replace with your template file path
    output_path = "/content/Output.docx"  # Replace with desired output path

    doc = Document(template_path)

    data = {
        "<<ProjectName>>": "Data Analyst Cover Letter Automation",
        "<<ProjectDescription>>": "The \"Data Analyst Cover Letter Automation\" project is a Python-based solution designed to streamline the process of generating personalized cover letters for the Data Analyst position. By harnessing the power of automation and leveraging the python-docx library, this project aims to enhance efficiency and accuracy in crafting compelling cover letters tailored to individual job applications.",
        "<<AuthorName>>": "Momin Naofil Ahmad",
        "<<AuthorEmail>>": "hashpro729@gmail.com",
        "<<AuthorContact>>": "9527338070",
        "<<StartDate>>": "August 1, 2023",
        "<<EndDate>>": "August 9, 2023",
    }

    for placeholder, value in data.items():
        replace_placeholder(doc, placeholder, value)

    doc.save(output_path)

if __name__ == "__main__":
    main()
